"""
er_core.py – Shared logic for ER Diagram Generator
ใช้ร่วมกันระหว่าง gen_er_diagram.py (CLI) และ app.py (Web UI)

แก้ไข v4:
  1. Dynamic table width  – คำนวณความกว้างจากชื่อ field ที่ยาวที่สุด
  2. Field name only      – ไม่แสดง data type ในเซลล์ field
  3. Thai name in header  – แสดงชื่อภาษาอังกฤษ + ภาษาไทยในหัวตาราง
  4. Left-align fields    – field ทุกเซลล์ชิดซ้าย
  5. Dynamic key column   – คอลัมน์ PK/FK/CHK ขยายตามความยาว key string
"""

import io
import re
import uuid
from collections import deque

try:
    from docx import Document
    from docx.oxml.ns import qn
    DOCX_OK = True
except ImportError:
    DOCX_OK = False

# ═══════════════════════════════════════════════════════════════════════════
# LAYOUT CONSTANTS
# ═══════════════════════════════════════════════════════════════════════════
# COL1_W ไม่ใช่ค่าคงที่อีกต่อไป — คำนวณ dynamic ต่อตาราง (_calc_col1_width)
MIN_COL1_W   = 30    # กว้างขั้นต่ำของคอลัมน์ key (px)
ROW_H        = 30
HEADER_H     = 30
GAP_W        = 60    # ช่องว่างแนวนอนระหว่างตาราง
GAP_H        = 50    # ช่องว่างแนวตั้งระหว่างแถว
START_X      = 40
START_Y      = 40
PAGE_W       = 1169  # A4 Landscape
PAGE_H       = 827
MAX_Y        = 755   # เว้น margin ล่าง
COLS_PER_ROW = 3

# ── Dynamic width ────────────────────────────────────────────────────────
MIN_TABLE_W   = 200   # กว้างขั้นต่ำของตาราง (px)
CHAR_PX       = 7.5   # pixel ต่อตัวอักษร 1 ตัว (font 12px Helvetica โดยประมาณ)
KEY_PADDING   = 14    # padding ซ้าย+ขวาของคอลัมน์ key
FIELD_PADDING = 22    # padding ฝั่งขวาของเซลล์ field
HDR_PADDING   = 24    # padding รวมของ header (ซ้าย+ขวา)
TYPE_PADDING  = 16    # padding ของคอลัมน์ datatype
MIN_TYPE_W    = 70    # กว้างขั้นต่ำของคอลัมน์ datatype (px)

# ═══════════════════════════════════════════════════════════════════════════
# DRAW.IO STYLES — คัดลอกจาก template 3_ER_จำหน่ายเครื่องจักรกล.drawio เป๊ะ 100%
# ═══════════════════════════════════════════════════════════════════════════
FONT_BASE     = "fontFamily=Helvetica;fontSize=12;"

STYLE_TABLE = (
    "shape=table;startSize=30;container=1;collapsible=1;childLayout=tableLayout;"
    "fixedRows=1;rowLines=0;fontStyle=1;align=center;resizeLast=1;"
    f"aspect=fixed;rounded=1;arcSize=15;{FONT_BASE}"
)
STYLE_ROW_PK = (
    "shape=tableRow;horizontal=0;startSize=0;swimlaneHead=0;swimlaneBody=0;"
    "fillColor=none;collapsible=0;dropTarget=0;points=[[0,0.5],[1,0.5]];"
    "portConstraint=eastwest;strokeColor=inherit;"
    "top=0;left=0;right=0;bottom=1;verticalAlign=middle;"
)
STYLE_ROW_NORMAL = (
    "shape=tableRow;horizontal=0;startSize=0;swimlaneHead=0;swimlaneBody=0;"
    "fillColor=none;collapsible=0;dropTarget=0;points=[[0,0.5],[1,0.5]];"
    "portConstraint=eastwest;strokeColor=inherit;"
    "top=0;left=0;right=0;bottom=0;verticalAlign=middle;"
)
STYLE_KEY_CELL = (
    "shape=partialRectangle;overflow=hidden;connectable=0;fillColor=none;"
    "strokeColor=inherit;top=0;left=0;bottom=0;right=0;"
    f"fontStyle=1;align=center;verticalAlign=middle;{FONT_BASE}"
)
STYLE_FIELD_PK = (
    "shape=partialRectangle;overflow=hidden;connectable=0;fillColor=none;"
    "align=left;strokeColor=inherit;top=0;left=0;bottom=0;right=0;"
    f"spacingLeft=6;fontStyle=5;verticalAlign=middle;{FONT_BASE}"
)
STYLE_FIELD_NORMAL = (
    "shape=partialRectangle;overflow=hidden;connectable=0;fillColor=none;"
    "align=left;strokeColor=inherit;top=0;left=0;bottom=0;right=0;"
    f"spacingLeft=6;verticalAlign=middle;{FONT_BASE}"
)
STYLE_TYPE_CELL = (
    "shape=partialRectangle;overflow=hidden;connectable=0;fillColor=none;"
    "align=left;strokeColor=inherit;top=0;left=0;bottom=0;right=0;"
    f"spacingLeft=6;verticalAlign=middle;{FONT_BASE}"
)
STYLE_TYPE_PK = (
    "shape=partialRectangle;overflow=hidden;connectable=0;fillColor=none;"
    "align=left;strokeColor=inherit;top=0;left=0;bottom=0;right=0;"
    f"spacingLeft=6;fontStyle=5;verticalAlign=middle;{FONT_BASE}"
)
# Crow's Foot edge
STYLE_EDGE = (
    "edgeStyle=entityRelationEdgeStyle;html=1;"
    "startArrow=ERmandOne;endArrow=ERzeroToMany;endFill=1;rounded=1;"
    "strokeColor=default;align=center;verticalAlign=middle;"
    f"fontColor=default;labelBackgroundColor=default;curved=0;{FONT_BASE}"
)

# ═══════════════════════════════════════════════════════════════════════════
# HELPERS
# ═══════════════════════════════════════════════════════════════════════════
def uid() -> str:
    return str(uuid.uuid4())

def is_pk(key_str: str) -> bool:
    return 'PK' in re.split(r'[,\s]+', key_str.upper())

def is_fk(key_str: str) -> bool:
    return bool(re.search(r'\bFK\d*\b', key_str.upper()))

def escape_xml(s: str) -> str:
    """Escape XML special chars + แปลง newline เป็น &#xa; สำหรับ XML attribute"""
    return (s.replace('&', '&amp;')
              .replace('<', '&lt;')
              .replace('>', '&gt;')
              .replace('"', '&quot;')
              .replace('\n', '&#xa;'))   # ← สำคัญสำหรับ 2-line header

def _extract_thai_name(description: str) -> str:
    """
    ดึงชื่อภาษาไทยออกจาก description
    Input:  "ตารางเอกสารสารบรรณ (DOH_DOCS)"
    Output: "ตารางเอกสารสารบรรณ"
    """
    thai = re.sub(r'\s*\([A-Z][A-Z0-9_]+\)\s*$', '', description).strip()
    return thai

def _calc_col3_width(tbl: dict) -> int:
    """คำนวณความกว้างคอลัมน์ datatype แบบ dynamic"""
    max_type = max((len(c['type']) for c in tbl['columns'] if c['type']), default=4)
    w = int(max_type * CHAR_PX) + TYPE_PADDING
    return max(MIN_TYPE_W, (w + 9) // 10 * 10)   # round up to nearest 10px

def _calc_col1_width(tbl: dict) -> int:
    """
    Fix 5: คำนวณความกว้างคอลัมน์ Key (PK/FK/CHK) แบบ dynamic
    - ใช้ key string ที่ยาวที่สุดในตารางนั้นเป็นตัวกำหนด
    - ตัวอย่าง: 'U1,U2,CHK1' (10 ตัว) → ต้องการ ~89px
    """
    max_key = max((len(c['key']) for c in tbl['columns'] if c['key']), default=2)
    w = int(max_key * CHAR_PX) + KEY_PADDING
    return max(MIN_COL1_W, (w + 4) // 5 * 5)   # round up to nearest 5px

def _calc_table_width(tbl: dict, col1_w: int, col3_w: int) -> int:
    """
    คำนวณความกว้างตารางอัตโนมัติ
    - col1_w: dynamic key column width
    - col3_w: dynamic datatype column width
    - ใช้ชื่อ field ที่ยาวที่สุดเป็นตัวกำหนด field column
    - คำนึงถึงความยาวชื่อตาราง (header) ด้วย
    """
    max_field = max((len(c['name']) for c in tbl['columns']), default=10)
    thai = _extract_thai_name(tbl.get('description', ''))
    max_hdr = max(len(tbl['name']), len(thai), 10)

    # field cell: col1_w + spacingLeft(6) + text + field padding + col3_w
    field_needed = col1_w + 6 + int(max_field * CHAR_PX) + FIELD_PADDING + col3_w
    # header (centered): text + padding ทั้งสองข้าง
    hdr_needed = int(max_hdr * CHAR_PX) + HDR_PADDING

    w = max(MIN_TABLE_W, field_needed, hdr_needed)
    return (w + 9) // 10 * 10   # round up to nearest 10px

# ═══════════════════════════════════════════════════════════════════════════
# STEP 1 — PARSE .docx
# ═══════════════════════════════════════════════════════════════════════════
def _get_cell_text(tc) -> str:
    parts = []
    for p in tc.findall(qn('w:p')):
        for t in p.iter(qn('w:t')):
            if t.text:
                parts.append(t.text)
    return ''.join(parts).strip()

def parse_docx(source) -> list[dict]:
    """
    รับ path (str/Path) หรือ bytes ของไฟล์ .docx
    คืน list ของ table dict:
      { "name": str, "description": str,
        "columns": [{"name","type","nullable","key","ref_table"}] }
    """
    if isinstance(source, bytes):
        doc = Document(io.BytesIO(source))
    else:
        doc = Document(source)

    tables = []
    current_heading = None

    for elem in doc.element.body:
        tag = elem.tag.split('}')[-1]

        if tag == 'p':
            pPr = elem.find(qn('w:pPr'))
            if pPr is not None:
                pStyle = pPr.find(qn('w:pStyle'))
                if pStyle is not None and re.match(
                        r'^(Heading\s*\d+|\d+)$',
                        pStyle.get(qn('w:val'), ''), re.IGNORECASE):
                    text = ''.join(
                        t.text for t in elem.iter(qn('w:t')) if t.text
                    ).strip()
                    m = re.search(r'\(([A-Z][A-Z0-9_]+)\)', text)
                    if m:
                        current_heading = (m.group(1), text)

        elif tag == 'tbl':
            if current_heading is None:
                continue   # intro/config table — skip

            tbl_name, tbl_desc = current_heading
            current_heading = None

            rows = elem.findall('.//' + qn('w:tr'))
            columns = []
            for row in rows[1:]:   # skip header row
                cells = row.findall('.//' + qn('w:tc'))
                if len(cells) < 7:
                    continue
                col_name  = _get_cell_text(cells[1])
                col_type  = _get_cell_text(cells[2])
                nullable  = _get_cell_text(cells[3])
                key_raw   = _get_cell_text(cells[5])
                ref_table = _get_cell_text(cells[6])

                if not col_name:
                    continue

                key_str = re.sub(r'\s+', '', key_raw)
                columns.append({
                    "name":      col_name,
                    "type":      col_type,
                    "nullable":  nullable.lower(),
                    "key":       key_str,
                    "ref_table": ref_table.strip(),
                })

            tables.append({
                "name":        tbl_name,
                "description": tbl_desc,
                "columns":     columns,
            })

    return tables

# ═══════════════════════════════════════════════════════════════════════════
# STEP 2 — FK-AWARE LAYOUT (A4 pagination)
# ═══════════════════════════════════════════════════════════════════════════

def _table_height(tbl: dict) -> int:
    return HEADER_H + ROW_H * len(tbl["columns"])

def _build_fk_graph(tables: list[dict]) -> dict[str, set]:
    table_names = {t['name'] for t in tables}
    adj: dict[str, set] = {t['name']: set() for t in tables}
    for tbl in tables:
        for col in tbl['columns']:
            ref = col['ref_table']
            if ref and ref in table_names:
                adj[tbl['name']].add(ref)
                adj[ref].add(tbl['name'])
    return adj

def _connected_components(tables: list[dict], adj: dict[str, set]) -> list[list[dict]]:
    table_map = {t['name']: t for t in tables}
    visited: set[str] = set()
    components: list[list[dict]] = []

    for tbl in tables:
        name = tbl['name']
        if name in visited:
            continue
        group: list[dict] = []
        q = deque([name])
        while q:
            n = q.popleft()
            if n in visited:
                continue
            visited.add(n)
            if n in table_map:
                group.append(table_map[n])
            for nb in sorted(adj.get(n, set())):
                if nb not in visited:
                    q.append(nb)
        components.append(group)

    return components

def layout_tables(tables: list[dict]) -> list[list[dict]]:
    """
    จัด layout A4 Landscape (3 col grid) พร้อม:
    - FK-aware grouping: ตารางที่เชื่อมกันอยู่หน้าเดียวกัน
    - Dynamic width: แต่ละตารางกว้างตามเนื้อหา
    - Column step = max(table width) + GAP_W เพื่อไม่ให้ทับกัน
    """
    # ── Pre-compute dynamic widths (col1 + col3 + total) ─────────────────
    tbl_col1   = {tbl['name']: _calc_col1_width(tbl)                                        for tbl in tables}
    tbl_col3   = {tbl['name']: _calc_col3_width(tbl)                                        for tbl in tables}
    tbl_widths = {tbl['name']: _calc_table_width(tbl, tbl_col1[tbl['name']], tbl_col3[tbl['name']]) for tbl in tables}
    max_tbl_w  = max(tbl_widths.values(), default=MIN_TABLE_W)
    col_step   = max_tbl_w + GAP_W   # ระยะ x ระหว่างคอลัมน์ (ใช้ตัวที่กว้างสุด)

    adj = _build_fk_graph(tables)
    components = _connected_components(tables, adj)

    pages: list[list[dict]] = []
    current_page: list[dict] = []
    row_x   = [START_X + col_step * c for c in range(COLS_PER_ROW)]
    col_idx = 0
    cur_y   = START_Y
    max_h_in_row = 0

    def flush_page():
        nonlocal current_page, col_idx, cur_y, max_h_in_row
        if current_page:
            pages.append(current_page)
        current_page = []
        col_idx = 0
        cur_y   = START_Y
        max_h_in_row = 0

    def add_table(tbl: dict):
        nonlocal col_idx, cur_y, max_h_in_row
        h = _table_height(tbl)
        if col_idx == COLS_PER_ROW:
            cur_y += max_h_in_row + GAP_H
            max_h_in_row = 0
            col_idx = 0
        t = dict(tbl)
        t['x']      = row_x[col_idx]
        t['y']      = cur_y
        t['width']  = tbl_widths[tbl['name']]   # ← dynamic total width
        t['col1_w'] = tbl_col1[tbl['name']]     # ← dynamic key-col width
        t['col3_w'] = tbl_col3[tbl['name']]     # ← dynamic type-col width
        t['height'] = h
        current_page.append(t)
        max_h_in_row = max(max_h_in_row, h)
        col_idx += 1

    def component_fits_on_new_row(comp: list[dict]) -> bool:
        rows_needed = (len(comp) + COLS_PER_ROW - 1) // COLS_PER_ROW
        max_h = max(
            max(_table_height(t) for t in comp[i:i + COLS_PER_ROW])
            for i in range(0, len(comp), COLS_PER_ROW)
        )
        offset = (max_h_in_row + GAP_H) if col_idx > 0 else 0
        return (cur_y + offset + rows_needed * (max_h + GAP_H)) <= MAX_Y

    for comp in components:
        if col_idx > 0 or current_page:
            if not component_fits_on_new_row(comp):
                if col_idx > 0:
                    cur_y += max_h_in_row + GAP_H
                    max_h_in_row = 0
                    col_idx = 0
                if cur_y + _table_height(comp[0]) > MAX_Y:
                    flush_page()

        for tbl in comp:
            h = _table_height(tbl)
            if col_idx == 0 and cur_y + h > MAX_Y and current_page:
                flush_page()
            add_table(tbl)

    if current_page:
        pages.append(current_page)

    return pages

# ═══════════════════════════════════════════════════════════════════════════
# STEP 3 — DRAW.IO XML GENERATION
# ═══════════════════════════════════════════════════════════════════════════

def _make_table_xml(tbl: dict) -> tuple[str, dict]:
    """สร้าง XML สำหรับ 1 ตาราง พร้อม fixes ทั้ง 5 ข้อ"""
    tbl_id  = uid()
    row_ids: dict[str, str] = {}
    x, y, w, h = tbl['x'], tbl['y'], tbl['width'], tbl['height']
    col1_w = tbl.get('col1_w', MIN_COL1_W)   # dynamic key-col width
    col3_w = tbl.get('col3_w', MIN_TYPE_W)   # dynamic type-col width
    col2_w = w - col1_w - col3_w             # field column = total - key - type

    # Fix 3: header = "EN_NAME\nชื่อภาษาไทย"
    thai_name  = _extract_thai_name(tbl.get('description', ''))
    tbl_label  = escape_xml(f"{tbl['name']}\n{thai_name}")

    lines = [
        f'<mxCell id="{tbl_id}" value="{tbl_label}" '
        f'style="{STYLE_TABLE}" vertex="1" parent="1">',
        f'  <mxGeometry x="{x}" y="{y}" width="{w}" height="{h}" as="geometry"/>',
        f'</mxCell>',
    ]

    row_y = 0
    for col in tbl['columns']:
        row_id = uid()
        row_ids[col['name']] = row_id

        pk = is_pk(col['key'])
        row_style   = STYLE_ROW_PK    if pk else STYLE_ROW_NORMAL
        field_style = STYLE_FIELD_PK  if pk else STYLE_FIELD_NORMAL

        # Row container
        lines += [
            f'<mxCell id="{row_id}" value="" style="{row_style}" '
            f'vertex="1" parent="{tbl_id}">',
            f'  <mxGeometry y="{HEADER_H + row_y}" width="{w}" height="{ROW_H}" as="geometry"/>',
            f'</mxCell>',
        ]

        # Fix 5: Key label cell — กว้าง dynamic col1_w
        lines += [
            f'<mxCell id="{uid()}" value="{escape_xml(col["key"])}" '
            f'style="{STYLE_KEY_CELL}" vertex="1" connectable="0" parent="{row_id}">',
            f'  <mxGeometry width="{col1_w}" height="{ROW_H}" as="geometry"/>',
            f'</mxCell>',
        ]

        # Field name cell (align=left)
        field_label = escape_xml(col['name'])
        lines += [
            f'<mxCell id="{uid()}" value="{field_label}" '
            f'style="{field_style}" vertex="1" connectable="0" parent="{row_id}">',
            f'  <mxGeometry x="{col1_w}" width="{col2_w}" height="{ROW_H}" as="geometry"/>',
            f'</mxCell>',
        ]

        # Datatype cell (left-align; PK row = bold+underline)
        type_style = STYLE_TYPE_PK if pk else STYLE_TYPE_CELL
        type_label = escape_xml(col.get('type', ''))
        lines += [
            f'<mxCell id="{uid()}" value="{type_label}" '
            f'style="{type_style}" vertex="1" connectable="0" parent="{row_id}">',
            f'  <mxGeometry x="{col1_w + col2_w}" width="{col3_w}" height="{ROW_H}" as="geometry"/>',
            f'</mxCell>',
        ]

        row_y += ROW_H

    return '\n'.join(lines), row_ids


def _make_edge_xml(src_row_id: str, tgt_row_id: str, label: str = "") -> str:
    """
    Crow's Foot edge:
      source = PK row ของตาราง parent (ฝั่ง "one"  → ERmandOne)
      target = FK row ของตาราง child  (ฝั่ง "many" → ERzeroToMany)
      label  = "fk_col = pk_col"
    """
    return (
        f'<mxCell id="{uid()}" value="{escape_xml(label)}" style="{STYLE_EDGE}" '
        f'edge="1" source="{src_row_id}" target="{tgt_row_id}" parent="1">'
        f'<mxGeometry relative="1" as="geometry"/>'
        f'</mxCell>'
    )


def _build_stub_tables(tables: list[dict]) -> list[dict]:
    """
    สร้าง stub table สำหรับตารางที่ถูกอ้างอิงผ่าน FK แต่ไม่มีในพจนานุกรม
    แสดง: PK (infer จาก FK column), ..., last_update_by, last_update_dtm
    """
    existing = {t['name'] for t in tables}
    refs: dict[str, list[str]] = {}
    for tbl in tables:
        for col in tbl['columns']:
            ref = col['ref_table']
            if ref and ref not in existing:
                refs.setdefault(ref, []).append(col['name'])

    stubs = []
    for ref_name, fk_cols in sorted(refs.items()):
        pk_col = fk_cols[0]   # infer PK จาก FK column แรกที่อ้างถึง
        stubs.append({
            "name":        ref_name,
            "description": "",   # ไม่มีคำอธิบายภาษาไทย
            "is_stub":     True,
            "columns": [
                {"name": pk_col, "type": "", "nullable": "", "key": "PK", "ref_table": ""},
                {"name": "...",  "type": "", "nullable": "", "key": "",   "ref_table": ""},
            ],
        })
    return stubs


def generate_drawio(tables: list[dict]) -> str:
    """Main function: tables → Draw.io XML string"""
    # รวม stub tables สำหรับ FK ที่ชี้ไปตารางนอก data dict
    stubs = _build_stub_tables(tables)
    all_tables = tables + stubs

    pages = layout_tables(all_tables)

    all_row_ids:   dict[str, dict] = {}
    page_of_table: dict[str, int]  = {}
    diagram_data:  list[list[str]] = []

    for page_idx, page_tables in enumerate(pages):
        cells: list[str] = []
        for tbl in page_tables:
            tbl_xml, row_id_map = _make_table_xml(tbl)
            cells.append(tbl_xml)
            all_row_ids[tbl['name']] = row_id_map
            page_of_table[tbl['name']] = page_idx
        diagram_data.append(cells)

    # ── Edges ─────────────────────────────────────────────────────────────
    edge_cells: dict[int, list[str]] = {i: [] for i in range(len(pages))}

    for tbl in all_tables:
        tbl_name = tbl['name']
        tbl_page = page_of_table.get(tbl_name)
        if tbl_page is None:
            continue
        tbl_row_ids = all_row_ids.get(tbl_name, {})

        for col in tbl['columns']:
            ref_table = col['ref_table']
            if not ref_table:
                continue
            ref_page = page_of_table.get(ref_table)
            if ref_page is None or ref_page != tbl_page:
                continue

            ref_tbl = next((t for t in all_tables if t['name'] == ref_table), None)
            if ref_tbl is None:
                continue

            pk_cols = [c for c in ref_tbl['columns'] if is_pk(c['key'])]
            if not pk_cols:
                continue

            src_id = all_row_ids.get(ref_table, {}).get(pk_cols[0]['name'])
            tgt_id = tbl_row_ids.get(col['name'])
            if src_id and tgt_id:
                label = f"{col['name']} = {pk_cols[0]['name']}"
                edge_cells[tbl_page].append(_make_edge_xml(src_id, tgt_id, label))

    # ── Assemble mxfile ────────────────────────────────────────────────────
    parts = [
        '<?xml version="1.0" encoding="UTF-8"?>',
        '<mxfile host="app.diagrams.net" type="device">',
    ]

    for page_idx, cells in enumerate(diagram_data):
        parts += [
            f'  <diagram name="Sheet{page_idx + 1}" id="{uid()}">',
            f'    <mxGraphModel dx="1620" dy="860" grid="1" gridSize="10" '
            f'guides="1" tooltips="1" connect="1" arrows="1" fold="1" '
            f'page="1" pageScale="1" pageWidth="{PAGE_W}" pageHeight="{PAGE_H}" '
            f'math="0" shadow="0">',
            '      <root>',
            '        <mxCell id="0"/>',
            '        <mxCell id="1" parent="0"/>',
        ]
        for cell_xml in cells:
            for line in cell_xml.splitlines():
                parts.append('        ' + line)
        for edge_xml in edge_cells[page_idx]:
            parts.append('        ' + edge_xml)
        parts += [
            '      </root>',
            '    </mxGraphModel>',
            '  </diagram>',
        ]

    parts.append('</mxfile>')
    return '\n'.join(parts)


# ═══════════════════════════════════════════════════════════════════════════
# STATS HELPER (สำหรับ Web UI)
# ═══════════════════════════════════════════════════════════════════════════
def get_stats(tables: list[dict], pages: list[list[dict]]) -> dict:
    page_of_table = {}
    for i, pg in enumerate(pages):
        for t in pg:
            page_of_table[t['name']] = i

    table_names = {t['name'] for t in tables}
    edge_count_per_page = [0] * len(pages)

    for tbl in tables:
        tbl_page = page_of_table.get(tbl['name'])
        if tbl_page is None:
            continue
        for col in tbl['columns']:
            ref = col['ref_table']
            if not ref or ref not in table_names:
                continue
            ref_page = page_of_table.get(ref)
            if ref_page is not None and ref_page == tbl_page:
                edge_count_per_page[tbl_page] += 1

    return {
        "table_count":  len(tables),
        "column_count": sum(len(t['columns']) for t in tables),
        "page_count":   len(pages),
        "pages": [
            {
                "tables":     [t['name'] for t in pg],
                "edge_count": edge_count_per_page[i],
            }
            for i, pg in enumerate(pages)
        ],
    }
